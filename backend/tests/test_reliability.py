"""
Phase 2 reliability tests — parser robustness + full pipeline with mocked LLM.

Tests:
  Parser (13 tests):
    - Normal TXT, PDF, DOCX
    - Corrupt PDF, Corrupt DOCX
    - Empty PDF, Empty TXT, Whitespace TXT
    - Old .doc format
    - Unsupported extension
    - Nonexistent file

  Pipeline with mocked LLM (7 tests):
    - Normal resume + JD
    - Resume with no experience / no skills / no education
    - JD with no experience requirement
    - JD with required + preferred skills
    - Minimal resume

  Error handling (4 tests):
    - Empty / whitespace resume or JD text

  Matcher edge cases (3 tests):
    - Empty resume vs full job
    - Overqualified candidate
    - No-requirements job
"""

from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
import PyPDF2
import docx

from app.services.parser import extract_text
from app.schemas.resume import ResumeProfile, Education, Experience
from app.schemas.job import JobProfile
from app.schemas.match import ScreeningOutput
from app.services.matcher import compute_match
from app.services.screener import screen_candidate


# ═══════════════════════════════════════════════════════════
# SAMPLE DATA
# ═══════════════════════════════════════════════════════════

SAMPLE_RESUME_TEXT = """\
JANE DOE
jane.doe@email.com | (555) 123-4567

EXPERIENCE
Senior Software Engineer — TechStart Inc.
June 2021 – Present
- Built REST APIs using FastAPI and PostgreSQL

Software Engineer — DataFlow Corp.
Jan 2019 – May 2021
- Developed ETL pipelines in Python

EDUCATION
B.S. Computer Science — State University, 2017

SKILLS
Python, FastAPI, PostgreSQL, Docker, AWS, Redis
"""

SAMPLE_JD_TEXT = """\
Senior Backend Engineer

Requirements:
- 5+ years of professional software development experience
- Strong proficiency in Python
- Experience with PostgreSQL
- Experience designing RESTful APIs
- Docker and Kubernetes

Nice to Have:
- Redis
- GraphQL
"""


# ═══════════════════════════════════════════════════════════
# FIXTURES
# ═══════════════════════════════════════════════════════════


@pytest.fixture
def txt_resume(tmp_path) -> Path:
    p = tmp_path / "resume.txt"
    p.write_text(SAMPLE_RESUME_TEXT, encoding="utf-8")
    return p


@pytest.fixture
def pdf_resume(tmp_path) -> Path:
    """Valid PDF with a page (text may not be extractable via PyPDF2 without reportlab,
    but the point is it doesn't crash)."""
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    pdf_path = tmp_path / "resume.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)
    return pdf_path


@pytest.fixture
def docx_resume(tmp_path) -> Path:
    doc = docx.Document()
    doc.add_heading("JANE DOE", level=1)
    doc.add_paragraph("jane.doe@email.com | (555) 123-4567")
    doc.add_heading("EXPERIENCE", level=2)
    doc.add_paragraph("Senior Software Engineer — TechStart Inc.")
    doc.add_heading("SKILLS", level=2)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker, AWS, Redis")
    doc.add_heading("EDUCATION", level=2)
    doc.add_paragraph("B.S. Computer Science — State University, 2017")

    docx_path = tmp_path / "resume.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def corrupt_pdf(tmp_path) -> Path:
    p = tmp_path / "corrupt.pdf"
    p.write_bytes(b"\x00\x01\x02\x03garbage data not a pdf")
    return p


@pytest.fixture
def empty_pdf(tmp_path) -> Path:
    from PyPDF2 import PdfWriter
    writer = PdfWriter()
    p = tmp_path / "empty.pdf"
    with open(p, "wb") as f:
        writer.write(f)
    return p


@pytest.fixture
def empty_txt(tmp_path) -> Path:
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    return p


@pytest.fixture
def whitespace_txt(tmp_path) -> Path:
    p = tmp_path / "whitespace.txt"
    p.write_text("   \n\n  \t  \n", encoding="utf-8")
    return p


@pytest.fixture
def corrupt_docx(tmp_path) -> Path:
    p = tmp_path / "corrupt.docx"
    p.write_bytes(b"this is not a docx file at all")
    return p


@pytest.fixture
def old_doc(tmp_path) -> Path:
    p = tmp_path / "old_format.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 100)
    return p


# ═══════════════════════════════════════════════════════════
# PARSER TESTS
# ═══════════════════════════════════════════════════════════


class TestParserTXT:
    def test_normal_txt(self, txt_resume):
        text = extract_text(txt_resume)
        assert len(text) > 50
        assert "JANE DOE" in text
        assert "Python" in text

    def test_empty_txt(self, empty_txt):
        text = extract_text(empty_txt)
        assert text == ""

    def test_whitespace_only_txt(self, whitespace_txt):
        text = extract_text(whitespace_txt)
        assert text.strip() == ""

    def test_sample_resume_txt(self):
        sample = Path(__file__).resolve().parent.parent.parent / "sample_data" / "sample_resume.txt"
        if sample.exists():
            text = extract_text(sample)
            assert "JANE DOE" in text
            assert "Python" in text


class TestParserPDF:
    def test_normal_pdf(self, pdf_resume):
        text = extract_text(pdf_resume)
        assert isinstance(text, str)
        # Blank page may yield empty string — the key is it doesn't crash

    def test_corrupt_pdf_doesnt_crash(self, corrupt_pdf):
        text = extract_text(corrupt_pdf)
        assert text == ""

    def test_empty_pdf(self, empty_pdf):
        text = extract_text(empty_pdf)
        assert text == ""


class TestParserDOCX:
    def test_normal_docx(self, docx_resume):
        text = extract_text(docx_resume)
        assert len(text) > 50
        assert "JANE DOE" in text
        assert "Python" in text

    def test_corrupt_docx_doesnt_crash(self, corrupt_docx):
        text = extract_text(corrupt_docx)
        assert text == ""


class TestParserEdgeCases:
    def test_old_doc_format(self, old_doc):
        text = extract_text(old_doc)
        assert text == ""

    def test_unsupported_extension(self, tmp_path):
        p = tmp_path / "file.xlsx"
        p.write_bytes(b"fake excel")
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(p)

    def test_nonexistent_file(self, tmp_path):
        p = tmp_path / "ghost.txt"
        text = extract_text(p)
        assert text == ""


class TestUploadValidation:
    """Test save_validated_upload helper for file type, size, and empty file validation."""

    @pytest.mark.anyio
    async def test_unsupported_file_extension(self):
        from fastapi import UploadFile, HTTPException
        from app.api.v1.router import save_validated_upload
        import io

        fake_file = UploadFile(filename="resume.png", file=io.BytesIO(b"png data"))
        with pytest.raises(HTTPException) as exc_info:
            await save_validated_upload(fake_file)
        assert exc_info.value.status_code == 400
        assert "Unsupported file type" in exc_info.value.detail

    @pytest.mark.anyio
    async def test_empty_file_upload(self):
        from fastapi import UploadFile, HTTPException
        from app.api.v1.router import save_validated_upload
        import io

        fake_file = UploadFile(filename="empty.pdf", file=io.BytesIO(b""))
        with pytest.raises(HTTPException) as exc_info:
            await save_validated_upload(fake_file)
        assert exc_info.value.status_code == 400
        assert "empty" in exc_info.value.detail

    @pytest.mark.anyio
    async def test_oversized_file_upload(self, monkeypatch):
        from fastapi import UploadFile, HTTPException
        from app.api.v1.router import save_validated_upload
        from app.core import config
        import io

        monkeypatch.setattr(config.settings, "max_upload_size_mb", 0.0001)

        fake_file = UploadFile(filename="large.txt", file=io.BytesIO(b"A" * 500))
        with pytest.raises(HTTPException) as exc_info:
            await save_validated_upload(fake_file)
        assert exc_info.value.status_code == 413
        assert "exceeds maximum allowed size" in exc_info.value.detail

    @pytest.mark.anyio
    async def test_valid_upload_saves_sanitized_uuid_file(self, tmp_path, monkeypatch):
        from fastapi import UploadFile
        from app.api.v1.router import save_validated_upload
        from app.core import config
        import io

        monkeypatch.setattr(config.settings, "upload_dir", str(tmp_path))

        fake_file = UploadFile(filename="My Resume (2026).pdf", file=io.BytesIO(b"PDF header content"))
        saved_path = await save_validated_upload(fake_file)

        assert saved_path.exists()
        assert saved_path.suffix == ".pdf"
        assert saved_path.name != "My Resume (2026).pdf"
        assert len(saved_path.name) > 30


# ═══════════════════════════════════════════════════════════
# MOCK LLM RESPONSES
# ═══════════════════════════════════════════════════════════

_MOCK_RESUME_PROFILE = {
    "name": "Jane Doe",
    "email": "jane.doe@email.com",
    "phone": "(555) 123-4567",
    "skills": ["Python", "FastAPI", "PostgreSQL", "Docker", "AWS", "Redis"],
    "education": [
        {"degree": "B.S.", "field": "Computer Science", "institution": "State University", "graduation_year": 2017}
    ],
    "experience": [
        {"company": "TechStart Inc.", "role": "Senior Software Engineer", "duration_months": 38, "description": "Built REST APIs"},
        {"company": "DataFlow Corp.", "role": "Software Engineer", "duration_months": 28, "description": "ETL pipelines"},
    ],
    "projects": [],
    "certifications": [],
}

_MOCK_JOB_PROFILE = {
    "job_title": "Senior Backend Engineer",
    "required_skills": ["Python", "PostgreSQL", "Docker", "Kubernetes", "REST APIs"],
    "preferred_skills": ["Redis", "GraphQL"],
    "experience_required": 5,
    "education_required": "Bachelor's degree",
    "responsibilities": ["Design APIs", "Scale services"],
}

_MOCK_REASONING = {
    "semantic_score": 82,
    "recommendation": "SHORTLIST",
    "reasoning": "Strong backend candidate with relevant experience.",
    "strengths": ["Solid Python and API experience", "Cloud experience with AWS"],
    "gaps": ["No Kubernetes experience mentioned"],
}


# ═══════════════════════════════════════════════════════════
# PIPELINE TESTS (mocked LLM)
# ═══════════════════════════════════════════════════════════


class TestPipelineNormal:
    @pytest.mark.anyio
    async def test_normal_screening(self):
        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [_MOCK_RESUME_PROFILE, _MOCK_JOB_PROFILE, _MOCK_REASONING]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)

            assert isinstance(output, ScreeningOutput)
            assert output.candidate_name == "Jane Doe"
            assert output.candidate_email == "jane.doe@email.com"
            assert output.match.final_score > 0
            assert output.match.skill_score > 0
            assert output.reasoning.recommendation == "SHORTLIST"
            assert len(output.reasoning.strengths) > 0


class TestPipelineEdgeCases:
    @pytest.mark.anyio
    async def test_resume_no_experience(self):
        no_exp_profile = {**_MOCK_RESUME_PROFILE, "experience": []}
        reasoning = {**_MOCK_REASONING, "semantic_score": 30, "recommendation": "REJECT"}

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [no_exp_profile, _MOCK_JOB_PROFILE, reasoning]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)
            assert output.match.experience_score < 10
            assert output.match.total_experience_months == 0

    @pytest.mark.anyio
    async def test_resume_no_skills(self):
        no_skills = {**_MOCK_RESUME_PROFILE, "skills": []}
        reasoning = {**_MOCK_REASONING, "semantic_score": 40, "recommendation": "REVIEW"}

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [no_skills, _MOCK_JOB_PROFILE, reasoning]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)
            assert output.match.skill_score < 30
            assert len(output.match.skill_details.missing_required) > 0

    @pytest.mark.anyio
    async def test_jd_no_experience_requirement(self):
        jd = {**_MOCK_JOB_PROFILE, "experience_required": None}

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [_MOCK_RESUME_PROFILE, jd, _MOCK_REASONING]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)
            assert output.match.experience_score == 100.0

    @pytest.mark.anyio
    async def test_jd_required_and_preferred_skills(self):
        jd = {**_MOCK_JOB_PROFILE,
              "required_skills": ["Python", "PostgreSQL", "Java"],
              "preferred_skills": ["Redis", "Kafka", "GraphQL"]}

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [_MOCK_RESUME_PROFILE, jd, _MOCK_REASONING]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)
            details = output.match.skill_details
            assert "Java" in details.missing_required
            assert len(details.matched_required) == 2  # Python, PostgreSQL
            assert len(details.matched_preferred) >= 1  # Redis

    @pytest.mark.anyio
    async def test_resume_no_education(self):
        no_edu = {**_MOCK_RESUME_PROFILE, "education": []}

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [no_edu, _MOCK_JOB_PROFILE, _MOCK_REASONING]

            output = await screen_candidate(SAMPLE_RESUME_TEXT, SAMPLE_JD_TEXT)
            assert isinstance(output, ScreeningOutput)
            assert output.match.education_score < 50

    @pytest.mark.anyio
    async def test_minimal_resume(self):
        minimal = {
            "name": "Nobody", "email": None, "phone": None,
            "skills": [], "education": [], "experience": [],
            "projects": [], "certifications": [],
        }
        reasoning = {
            "semantic_score": 5, "recommendation": "REJECT",
            "reasoning": "No relevant experience.", "strengths": [],
            "gaps": ["No skills", "No experience"],
        }

        with patch("app.services.llm._call_openai", new_callable=AsyncMock) as mock_llm, \
             patch("app.services.llm.settings") as mock_settings:
            mock_settings.llm_provider = "openai"
            mock_llm.side_effect = [minimal, _MOCK_JOB_PROFILE, reasoning]

            output = await screen_candidate("Nobody", SAMPLE_JD_TEXT)
            assert output.candidate_name == "Nobody"
            assert output.candidate_email is None
            assert output.match.final_score < 30


# ═══════════════════════════════════════════════════════════
# ERROR HANDLING TESTS
# ═══════════════════════════════════════════════════════════


class TestPipelineErrors:
    @pytest.mark.anyio
    async def test_empty_resume_text_raises(self):
        with pytest.raises(ValueError, match="Resume text is empty"):
            await screen_candidate("", SAMPLE_JD_TEXT)

    @pytest.mark.anyio
    async def test_whitespace_resume_text_raises(self):
        with pytest.raises(ValueError, match="Resume text is empty"):
            await screen_candidate("   \n  \t  ", SAMPLE_JD_TEXT)

    @pytest.mark.anyio
    async def test_empty_jd_text_raises(self):
        with pytest.raises(ValueError, match="Job description text is empty"):
            await screen_candidate(SAMPLE_RESUME_TEXT, "")

    @pytest.mark.anyio
    async def test_whitespace_jd_text_raises(self):
        with pytest.raises(ValueError, match="Job description text is empty"):
            await screen_candidate(SAMPLE_RESUME_TEXT, "   \n  ")


# ═══════════════════════════════════════════════════════════
# MATCHER EDGE CASES (deterministic, no mocking)
# ═══════════════════════════════════════════════════════════


class TestMatcherEdgeCases:
    def test_empty_resume_vs_full_job(self):
        resume = ResumeProfile()
        job = JobProfile(
            job_title="Engineer",
            required_skills=["Python", "Java"],
            preferred_skills=["Docker"],
            experience_required=3,
            education_required="Bachelor's",
        )
        result = compute_match(resume, job, semantic_score=0)
        assert result.final_score < 20
        assert len(result.gaps) > 0

    def test_overqualified_candidate(self):
        resume = ResumeProfile(
            skills=["Python", "Java", "C++", "Go", "Rust", "Docker", "K8s"],
            experience=[Experience(company="BigCo", role="Staff Engineer", duration_months=120)],
            education=[Education(degree="Ph.D.", field="Computer Science")],
        )
        job = JobProfile(
            job_title="Junior Dev",
            required_skills=["Python"],
            experience_required=1,
            education_required="Bachelor's",
        )
        result = compute_match(resume, job, semantic_score=90)
        assert result.final_score > 80

    def test_no_requirements_job(self):
        resume = ResumeProfile(name="Anyone", skills=["Excel"])
        job = JobProfile(job_title="Open Role")
        result = compute_match(resume, job, semantic_score=50)
        assert result.skill_score == 100.0
        assert result.experience_score == 100.0
        assert result.education_score == 100.0
